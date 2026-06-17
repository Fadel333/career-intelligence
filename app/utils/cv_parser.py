import os
import re
from typing import Dict, List, Any
from PyPDF2 import PdfReader
from docx import Document
from app.utils.advanced_cv_parser import AdvancedCVParser

class CVParser:
    """Main CV parser interface using advanced parsing"""
    
    @staticmethod
    def parse_cv(file_path: str) -> Dict[str, Any]:
        """Parse CV and return structured data"""
        # First extract text from file
        text = CVParser._extract_text(file_path)
        
        if not text:
            return None
        
        # Use advanced parser
        parser = AdvancedCVParser()
        
        # For large files, use quick parsing for speed
        if len(text) > 10000:
            results = parser.parse_cv_quick(text)
        else:
            results = parser.parse_cv(text)
        
        # Calculate total skills
        results['total_skills'] = sum(len(skills) for skills in results['skills'].values())
        
        return results
    
    @staticmethod
    def _extract_text(file_path: str) -> str:
        """Extract text from PDF or DOCX with optimized reading"""
        text = ""
        
        if file_path.endswith('.pdf'):
            try:
                reader = PdfReader(file_path)
                # Limit pages for speed (max 5 pages for quick parsing)
                max_pages = min(len(reader.pages), 5)
                for i in range(max_pages):
                    page_text = reader.pages[i].extract_text()
                    if page_text:
                        text += page_text + "\n"
            except Exception as e:
                print(f"Error reading PDF: {e}")
                
        elif file_path.endswith('.docx'):
            try:
                doc = Document(file_path)
                # Limit paragraphs for speed
                max_paragraphs = min(len(doc.paragraphs), 50)
                for i in range(max_paragraphs):
                    text += doc.paragraphs[i].text + "\n"
            except Exception as e:
                print(f"Error reading DOCX: {e}")
        
        return text
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF file (legacy method)"""
        return CVParser._extract_text(file_path)
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from DOCX file (legacy method)"""
        return CVParser._extract_text(file_path)