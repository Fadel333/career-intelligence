# app/utils/hybrid_parser.py
import time
import threading
from typing import Dict, Any, Optional
from app.utils.advanced_cv_parser import AdvancedCVParser
from app.utils.cv_detector import CVDetector
import os
from PyPDF2 import PdfReader
from docx import Document


class HybridParser:
    """Hybrid CV parser - Quick parse first, deep parse in background with CV detection"""
    
    def __init__(self):
        # Lazy load - only load when needed
        self._quick_parser = None
        self._deep_parser = None
        self.cache = {}
        self.processing_queue = {}
        
        # Configuration
        self.MAX_WORDS_FOR_QUICK_ONLY = 500
        self.DEEP_PARSE_DELAY = 2  # seconds
        self.CACHE_MAX_SIZE = 100  # Maximum number of cached items
        self.MIN_CV_CONFIDENCE = 30.0  # Minimum confidence to consider as CV (30%)
        
        print("🔧 HybridParser initialized with lazy loading and CV detection")
    
    def _get_quick_parser(self):
        """Lazy load quick parser"""
        if self._quick_parser is None:
            print("⏳ Loading quick parser...")
            try:
                # REMOVED the mode='quick' parameter
                self._quick_parser = AdvancedCVParser()
                print("✅ Quick parser loaded")
            except Exception as e:
                print(f"❌ Error loading quick parser: {e}")
                raise
        return self._quick_parser
    
    def _get_deep_parser(self):
        """Lazy load deep parser"""
        if self._deep_parser is None:
            print("⏳ Loading deep parser...")
            try:
                # REMOVED the mode='deep' parameter
                self._deep_parser = AdvancedCVParser()
                print("✅ Deep parser loaded")
            except Exception as e:
                print(f"❌ Error loading deep parser: {e}")
                raise
        return self._deep_parser
    
    def detect_cv(self, text: str) -> tuple:
        """
        Detect if the text is a CV using CVDetector
        
        Args:
            text: Extracted text from file
            
        Returns:
            Tuple of (is_cv, confidence, reason)
        """
        try:
            is_cv, confidence, reason = CVDetector.detect(text)
            
            # Get detailed report for debugging
            if confidence < 50:
                report = CVDetector.get_detailed_report(text)
                print(f"📊 CV Detection Report:")
                print(f"   - Headers found: {report['headers_count']}")
                print(f"   - Patterns matched: {report['patterns_count']}")
                print(f"   - Word count: {report['word_count']}")
                print(f"   - Has sections: {report['has_sections']}")
                print(f"   - Has bullets: {report['has_bullets']}")
                print(f"   - Has dates: {report['has_dates']}")
                print(f"   - Has email: {report['has_email']}")
                print(f"   - Has phone: {report['has_phone']}")
                print(f"   - Has degree: {report['has_degree']}")
            
            print(f"🔍 CV Detection: {reason} (Confidence: {confidence:.1f}%)")
            return is_cv, confidence, reason
        except Exception as e:
            print(f"❌ CV detection error: {e}")
            # Default to allowing parsing if detection fails
            return True, 50.0, "Detection failed, allowing parse"
    
    def parse_hybrid(self, file_path: str, user_id: int) -> Optional[Dict[str, Any]]:
        """
        Hybrid parsing: Detect CV first, then quick parse, deep parse in background
        
        Args:
            file_path: Path to the CV file
            user_id: ID of the user uploading the CV
            
        Returns:
            Dictionary with parsed data or None if parsing fails
        """
        
        print(f"🔍 HybridParser.parse_hybrid called for {file_path}")
        print(f"👤 User ID: {user_id}")
        
        # Check if file exists
        if not os.path.exists(file_path):
            print(f"❌ File not found: {file_path}")
            return None
        
        # Check cache first
        cache_key = f"{user_id}_{os.path.basename(file_path)}"
        if cache_key in self.cache:
            cached_data = self.cache[cache_key]
            print(f"📦 Returning cached data for {cache_key}")
            return cached_data
        
        # Step 1: Extract text (optimized)
        print("📄 Extracting text from file...")
        text = self._extract_text_fast(file_path)
        if not text:
            print("❌ No text extracted from file")
            return None
        
        word_count = len(text.split())
        print(f"📄 Extracted {len(text)} characters, {word_count} words")
        
        # Step 2: Detect if it's a CV
        print("🔍 Running CV detection...")
        is_cv, confidence, reason = self.detect_cv(text)
        
        # Get detailed report for logging
        report = CVDetector.get_detailed_report(text)
        
        if not is_cv:
            print(f"❌ Not a CV: {reason} (Confidence: {confidence:.1f}%)")
            return {
                'is_cv': False,
                'confidence': confidence,
                'reason': reason,
                'status': 'rejected',
                'error': f"File is not a valid CV: {reason}",
                'detection_report': report
            }
        
        if confidence < self.MIN_CV_CONFIDENCE:
            print(f"⚠️ Low confidence CV: {confidence:.1f}% < {self.MIN_CV_CONFIDENCE}%")
            cv_warning = f"Low confidence CV detection ({confidence:.1f}%)"
        else:
            cv_warning = None
        
        print(f"✅ Valid CV detected with {confidence:.1f}% confidence")
        print(f"📊 Detection details: {report['headers_count']} headers, {report['patterns_count']} patterns")
        
        # Step 3: Quick parse
        print("⚡ Running quick parse...")
        quick_results = self._quick_parse(text)
        print(f"✅ Quick parse complete: {quick_results.get('total_skills', 0)} skills found")
        print(f"📂 Skills categories: {list(quick_results.get('skills', {}).keys())}")
        
        # Add CV detection info to results
        quick_results['is_cv'] = True
        quick_results['cv_confidence'] = confidence
        quick_results['cv_reason'] = reason
        quick_results['cv_warning'] = cv_warning
        quick_results['detection_report'] = report
        
        # If CV is small, skip deep parse
        if word_count < self.MAX_WORDS_FOR_QUICK_ONLY:
            print(f"📄 Small CV ({word_count} words < {self.MAX_WORDS_FOR_QUICK_ONLY}) - skipping deep parse")
            quick_results['status'] = 'complete'
            quick_results['deep_parse_started'] = False
            quick_results['parse_type'] = 'quick_only'
            quick_results['word_count'] = word_count
            
            # Cache the result
            self._add_to_cache(cache_key, quick_results)
            return quick_results
        
        # Step 4: Start deep parse in background thread for larger CVs
        self._start_deep_parse(text, quick_results, cache_key, user_id)
        
        # Step 5: Return quick results immediately
        quick_results['status'] = 'processing'
        quick_results['deep_parse_started'] = True
        quick_results['parse_type'] = 'quick_initial'
        quick_results['word_count'] = word_count
        
        # Cache the quick results
        self._add_to_cache(cache_key, quick_results)
        
        print(f"✅ Returning quick results for {cache_key}")
        print(f"📊 Skills found: {quick_results.get('total_skills', 0)}")
        return quick_results
    
    def _quick_parse(self, text: str) -> Dict[str, Any]:
        """
        Quick parse - fast but less accurate (1-2 seconds)
        
        Args:
            text: Extracted text from CV
            
        Returns:
            Dictionary with parsed data
        """
        start_time = time.time()
        
        try:
            # Use quick parsing method
            parser = self._get_quick_parser()
            results = parser.parse_cv_quick(text)
            
            # Ensure the results have the expected structure
            if 'skills' not in results:
                results['skills'] = {}
            
            if 'total_skills' not in results:
                total = 0
                for category, skills in results.get('skills', {}).items():
                    if isinstance(skills, list):
                        total += len(skills)
                    elif isinstance(skills, (str, int, float)):
                        total += 1
                results['total_skills'] = total
            
            # Add timing info
            results['parse_time'] = round(time.time() - start_time, 2)
            results['parse_type'] = 'quick'
            results['status'] = 'processing'
            
            print(f"⚡ Quick parse completed in {results['parse_time']}s")
            print(f"📊 Found {results['total_skills']} skills")
            print(f"📂 Categories: {list(results['skills'].keys())}")
            
            return results
            
        except Exception as e:
            print(f"❌ Quick parse error: {e}")
            import traceback
            traceback.print_exc()
            
            # Return minimal results
            return {
                'skills': {},
                'total_skills': 0,
                'experience_years': 0,
                'education': [],
                'certifications': [],
                'parse_time': round(time.time() - start_time, 2),
                'parse_type': 'quick',
                'status': 'error',
                'error': str(e)
            }
    
    def _deep_parse(self, text: str, quick_results: Dict, cache_key: str) -> Dict:
        """
        Deep parse - slow but accurate (runs in background)
        
        Args:
            text: Extracted text from CV
            quick_results: Results from quick parse
            cache_key: Cache key for storing results
            
        Returns:
            Merged results with deep parse data
        """
        start_time = time.time()
        
        try:
            print(f"🔍 Starting deep parse for {cache_key}...")
            
            # Use full parsing
            parser = self._get_deep_parser()
            deep_results = parser.parse_cv(text)
            
            # Merge results (deep results override quick results)
            merged = {**quick_results, **deep_results}
            
            # Ensure skills are merged properly
            if 'skills' in deep_results and isinstance(deep_results['skills'], dict):
                if 'skills' in quick_results and isinstance(quick_results['skills'], dict):
                    # Deep merge skills
                    for category, skills in deep_results['skills'].items():
                        if category in quick_results['skills']:
                            # Combine and deduplicate
                            existing = quick_results['skills'].get(category, [])
                            if isinstance(existing, list) and isinstance(skills, list):
                                merged['skills'][category] = list(set(existing + skills))
                            elif isinstance(skills, list):
                                merged['skills'][category] = skills
                        else:
                            merged['skills'][category] = skills
            
            # Recalculate total skills
            total = 0
            for category, skills in merged.get('skills', {}).items():
                if isinstance(skills, list):
                    total += len(skills)
                elif isinstance(skills, (str, int, float)):
                    total += 1
            merged['total_skills'] = total
            
            # Add timing info
            merged['parse_time'] = round(time.time() - start_time, 2)
            merged['parse_type'] = 'deep'
            merged['status'] = 'complete'
            merged['deep_parse_started'] = True
            
            # Preserve CV detection info
            if 'is_cv' in quick_results:
                merged['is_cv'] = quick_results['is_cv']
                merged['cv_confidence'] = quick_results.get('cv_confidence', 0)
                merged['cv_reason'] = quick_results.get('cv_reason', '')
                merged['cv_warning'] = quick_results.get('cv_warning', None)
                merged['detection_report'] = quick_results.get('detection_report', {})
            
            # Update cache
            self._add_to_cache(cache_key, merged)
            self.processing_queue.pop(cache_key, None)
            
            print(f"✅ Deep parse complete for {cache_key} in {merged['parse_time']}s")
            print(f"📊 Total skills: {merged['total_skills']}")
            print(f"📂 Categories: {list(merged['skills'].keys())}")
            return merged
            
        except Exception as e:
            print(f"❌ Deep parse error for {cache_key}: {e}")
            import traceback
            traceback.print_exc()
            
            self.processing_queue.pop(cache_key, None)
            
            # Return quick results as fallback
            quick_results['status'] = 'complete'
            quick_results['deep_parse_started'] = False
            quick_results['parse_error'] = str(e)
            quick_results['parse_type'] = 'quick_fallback'
            
            # Preserve CV detection info
            if 'is_cv' in quick_results:
                quick_results['cv_detection'] = {
                    'is_cv': quick_results['is_cv'],
                    'confidence': quick_results.get('cv_confidence', 0),
                    'reason': quick_results.get('cv_reason', ''),
                    'warning': quick_results.get('cv_warning', None)
                }
            
            self._add_to_cache(cache_key, quick_results)
            return quick_results
    
    def _start_deep_parse(self, text: str, quick_results: Dict, cache_key: str, user_id: int):
        """
        Start deep parse in background thread
        
        Args:
            text: Extracted text from CV
            quick_results: Results from quick parse
            cache_key: Cache key for storing results
            user_id: ID of the user
        """
        
        # Prevent duplicate processing
        if cache_key in self.processing_queue:
            print(f"⏳ {cache_key} already processing")
            return
        
        self.processing_queue[cache_key] = {
            'started_at': time.time(),
            'user_id': user_id,
            'status': 'processing'
        }
        
        def parse_in_background():
            try:
                # Wait a bit before starting deep parse (let user see quick results first)
                time.sleep(self.DEEP_PARSE_DELAY)
                self._deep_parse(text, quick_results, cache_key)
            except Exception as e:
                print(f"❌ Background parse error: {e}")
                import traceback
                traceback.print_exc()
                self.processing_queue.pop(cache_key, None)
        
        thread = threading.Thread(target=parse_in_background, daemon=True)
        thread.start()
        print(f"🔄 Background deep parse started for {cache_key}")
        print(f"⏱️  Deep parse will run in {self.DEEP_PARSE_DELAY}s")
    
    def _extract_text_fast(self, file_path: str) -> str:
        """
        Extract text from file - optimized for speed
        
        Args:
            file_path: Path to the file
            
        Returns:
            Extracted text as string
        """
        text = ""
        
        try:
            if file_path.endswith('.pdf'):
                print(f"📄 Extracting from PDF: {file_path}")
                reader = PdfReader(file_path)
                # Only read first 3 pages for quick parse
                max_pages = min(len(reader.pages), 3)
                print(f"📄 Reading {max_pages} pages")
                
                for i in range(max_pages):
                    try:
                        page_text = reader.pages[i].extract_text()
                        if page_text:
                            text += page_text + "\n"
                    except Exception as e:
                        print(f"⚠️ Error reading page {i}: {e}")
                        continue
                        
            elif file_path.endswith('.docx'):
                print(f"📄 Extracting from DOCX: {file_path}")
                doc = Document(file_path)
                # Only read first 30 paragraphs
                max_paras = min(len(doc.paragraphs), 30)
                print(f"📄 Reading {max_paras} paragraphs")
                
                for i in range(max_paras):
                    try:
                        if doc.paragraphs[i].text:
                            text += doc.paragraphs[i].text + "\n"
                    except Exception as e:
                        print(f"⚠️ Error reading paragraph {i}: {e}")
                        continue
            else:
                print(f"❌ Unsupported file type: {file_path}")
                return ""
                
        except Exception as e:
            print(f"❌ Error extracting text: {e}")
            import traceback
            traceback.print_exc()
            return ""
        
        print(f"📄 Extracted {len(text)} characters")
        return text.strip()
    
    def _add_to_cache(self, cache_key: str, data: Dict):
        """
        Add data to cache with size management
        
        Args:
            cache_key: Cache key
            data: Data to cache
        """
        # If cache is too large, remove oldest entries
        if len(self.cache) >= self.CACHE_MAX_SIZE:
            # Remove 20% of oldest entries
            keys_to_remove = list(self.cache.keys())[:int(self.CACHE_MAX_SIZE * 0.2)]
            for key in keys_to_remove:
                del self.cache[key]
            print(f"🧹 Cache cleanup: removed {len(keys_to_remove)} entries")
        
        self.cache[cache_key] = data
        print(f"💾 Added to cache: {cache_key}")
    
    def get_parsed_data(self, user_id: int, file_path: str) -> Optional[Dict]:
        """
        Get parsed data from cache
        
        Args:
            user_id: ID of the user
            file_path: Path to the file
            
        Returns:
            Cached data or None if not found
        """
        cache_key = f"{user_id}_{os.path.basename(file_path)}"
        return self.cache.get(cache_key)
    
    def is_parsing_complete(self, user_id: int, file_path: str) -> bool:
        """
        Check if deep parsing is complete
        
        Args:
            user_id: ID of the user
            file_path: Path to the file
            
        Returns:
            True if parsing is complete, False otherwise
        """
        cache_key = f"{user_id}_{os.path.basename(file_path)}"
        data = self.cache.get(cache_key)
        if data:
            return data.get('status') == 'complete'
        return False
    
    def get_parse_status(self, user_id: int, file_path: str) -> Dict:
        """
        Get detailed parse status including CV detection info
        
        Args:
            user_id: ID of the user
            file_path: Path to the file
            
        Returns:
            Dictionary with status information
        """
        cache_key = f"{user_id}_{os.path.basename(file_path)}"
        data = self.cache.get(cache_key)
        
        if data:
            status_info = {
                'status': data.get('status', 'unknown'),
                'parse_type': data.get('parse_type', 'unknown'),
                'total_skills': data.get('total_skills', 0),
                'parse_time': data.get('parse_time', 0),
                'deep_parse_started': data.get('deep_parse_started', False),
                'word_count': data.get('word_count', 0),
                'error': data.get('error', None)
            }
            
            # Add CV detection info if available
            if 'is_cv' in data:
                status_info['cv_detection'] = {
                    'is_cv': data.get('is_cv', False),
                    'confidence': data.get('cv_confidence', 0),
                    'reason': data.get('cv_reason', ''),
                    'warning': data.get('cv_warning', None)
                }
            
            return status_info
        
        # Check if in processing queue
        if cache_key in self.processing_queue:
            return {
                'status': 'processing',
                'parse_type': 'quick_initial',
                'total_skills': 0,
                'parse_time': 0,
                'deep_parse_started': True,
                'word_count': 0,
                'error': None,
                'cv_detection': None
            }
        
        return {'status': 'not_found'}
    
    def clear_cache(self, user_id: Optional[int] = None):
        """
        Clear cache for a user or all users
        
        Args:
            user_id: Optional user ID to clear cache for specific user
        """
        if user_id is not None:
            keys_to_delete = [k for k in self.cache.keys() if k.startswith(f"{user_id}_")]
            for key in keys_to_delete:
                del self.cache[key]
            print(f"🧹 Cleared cache for user {user_id} ({len(keys_to_delete)} entries)")
        else:
            cache_size = len(self.cache)
            self.cache.clear()
            self.processing_queue.clear()
            print(f"🧹 Cleared all cache ({cache_size} entries)")
    
    def get_cache_stats(self) -> Dict:
        """
        Get cache statistics
        
        Returns:
            Dictionary with cache statistics
        """
        return {
            'cache_size': len(self.cache),
            'processing_queue': len(self.processing_queue),
            'cache_keys': list(self.cache.keys()),
            'processing_keys': list(self.processing_queue.keys())
        }
    
    def validate_cv_file(self, file_path: str) -> Dict:
        """
        Validate if a file is a CV without full parsing
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary with validation results
        """
        print(f"🔍 Validating CV file: {file_path}")
        
        # Check if file exists
        if not os.path.exists(file_path):
            return {
                'valid': False,
                'error': 'File does not exist',
                'is_cv': False,
                'confidence': 0,
                'reason': 'File not found'
            }
        
        # Extract text
        text = self._extract_text_fast(file_path)
        if not text:
            return {
                'valid': False,
                'error': 'No text extracted from file',
                'is_cv': False,
                'confidence': 0,
                'reason': 'Empty file or unsupported format'
            }
        
        # Detect if it's a CV
        is_cv, confidence, reason = self.detect_cv(text)
        
        # Get detailed report
        report = CVDetector.get_detailed_report(text)
        
        return {
            'valid': is_cv,
            'is_cv': is_cv,
            'confidence': confidence,
            'reason': reason,
            'word_count': len(text.split()),
            'character_count': len(text),
            'detection_report': report
        }